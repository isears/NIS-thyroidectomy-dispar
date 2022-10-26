import pandas as pd
import numpy as np
from nistd.dataProcessing import categorical_lookup
import docx


def make_cell_bold(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True


if __name__ == "__main__":
    processed_df = pd.read_csv("cache/preprocessed.csv")

    # TODO: oooof
    categorical_lookup["INCOME_QRTL"] = list(range(1, 5))

    doc = docx.Document()
    table_doc = doc.add_table(rows=1, cols=2)
    table_doc.autofit = True
    table_doc.allow_autofit = True

    header_row = table_doc.row_cells(0)
    header_row[0].text = "Demographics and Preoperative Characteristics"
    header_row[1].text = "N (%)"

    categoricals = processed_df[
        [c for c in processed_df.columns if c in categorical_lookup]
    ]

    table1_df = pd.DataFrame()

    # Get age mean / sem
    age_mean = processed_df["AGE"].mean()
    age_sem = processed_df["AGE"].std() / np.sqrt(len(processed_df))
    as_str = f"{age_mean:.2f} +/- {age_sem:.2f}"
    table1_df = pd.concat(
        [table1_df, pd.DataFrame(data={"N (%)": as_str}, index=["Age mean"])]
    )

    row = table_doc.add_row().cells
    row[0].text = "Age"
    make_cell_bold(row[0])
    row[1].text = as_str

    # Get > 65
    over65_count = (processed_df["AGE"] > 65).sum()
    as_str = f"{over65_count:,} ({100 * over65_count / len(processed_df):.2f})"
    table1_df = pd.concat(
        [table1_df, pd.DataFrame(data={"N (%)": as_str}, index=["AGE > 65"])]
    )

    row = table_doc.add_row().cells
    row[0].text = ">65"
    row[1].text = as_str

    # Get APDRG stats
    row = table_doc.add_row().cells
    row[0].text = "APDRG"
    make_cell_bold(row[0])

    for aprdrg_col in ["APRDRG_Severity", "APRDRG_Risk_Mortality"]:
        over2_count = (processed_df[aprdrg_col] > 2).sum()
        as_str = f"{over2_count:,} ({100 * over2_count / len(processed_df):.2f})"
        table1_df = pd.concat(
            [
                table1_df,
                pd.DataFrame(data={"N (%)": as_str}, index=[f"{aprdrg_col} > 2"]),
            ]
        )

        row = table_doc.add_row().cells
        row[0].text = f"{aprdrg_col} > 2"
        row[1].text = as_str

    for variable_name in categoricals.columns:
        row = table_doc.add_row().cells
        row[0].text = variable_name
        make_cell_bold(row[0])

        vc = categoricals[variable_name].value_counts()

        def format_n(n):
            percent = n / len(processed_df) * 100
            return f"{n:,} ({percent:.2f})"

        this_var_df = pd.DataFrame(
            data={"N (%)": vc.apply(format_n).values},
            index=vc.index.map(lambda x: f"[{vc.name}] {x}"),
        )

        table1_df = table1_df.append(this_var_df)

        for subcat, n in zip(vc.index.to_list(), vc.to_list()):
            row = table_doc.add_row().cells
            row[0].text = str(subcat)
            row[1].text = format_n(n)

    table1_df.to_csv("results/table1.csv")

    table_doc.style = "Table Grid"

    doc.save(f"results/table1.docx")
